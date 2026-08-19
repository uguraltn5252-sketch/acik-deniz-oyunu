from pathlib import Path
from hashlib import sha256
from docx import Document
import sys

ROOT=Path(__file__).resolve().parent
RULE=ROOT/'OYUN_Kural_Kitabi_v2.6.docx'
CARDS=ROOT/'OYUN_Kartlar_A4_Prototip_v2.5_UNCHANGED.pdf'
CHARS=ROOT/'OYUN_Karakter_Kartlari_v2.5_UNCHANGED.pdf'
MOD=ROOT/'OYUN_Moderator_Masa_Karti_v2.6.pdf'
BASE=ROOT/'MEKANIK_BASELINE_OYUN_SIMULASYON_PAKETI_v2.5.zip'

expected_card_hash='e158b33b77d2fff962420170d87aea407c87c97c9d611e19a6b72e7827aba4cc'
expected_base_hash='975dd77d435a835fcf3faa864c5624b0542fadb419fa7d34217ed32fe87aa046'

def h(p):
    x=sha256()
    with p.open('rb') as f:
        for b in iter(lambda:f.read(1024*1024),b''): x.update(b)
    return x.hexdigest()

def all_text(doc):
    out=[]
    for p in doc.paragraphs: out.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            out.extend(c.text for c in r.cells)
    return '\n'.join(out)

fail=[]
for p in [RULE,CARDS,CHARS,MOD,BASE]:
    if not p.exists(): fail.append(f'missing {p.name}')
if fail:
    print('FAIL',fail); sys.exit(1)

if h(CARDS)!=expected_card_hash: fail.append('full card PDF hash changed')
if h(BASE)!=expected_base_hash: fail.append('v2.5 baseline ZIP hash changed')

d=Document(RULE); txt=all_text(d)
checks={
'player_count':'6-15 oyuncu + 1 tarafsız Moderatör',
'hull2':'Gemi 2 Gövdeyle başlar',
'captain_before_loyalty':'Kaptanı Sadakatler dağıtılmadan önce seçtiniz',
'captain_vote2':'Kaptanın rota oyu 2 sayar',
'first_night_one':'tam 1 kartın olay yüzüne gizlice bakar',
'first_traitor_no_attack':'Bu ilk Hain gecesinde saldırı yoktur',
'scurvy_island':'Liman Gecesinden önce herhangi bir Ada kartına girmek zorundasınız',
'open_stays_open':'Kamusal açılmış fakat ziyaret edilmemiş kartın bilgisi açıktır',
'captain_replacement':'Kaptan ölür, Kamaraya girer, mahsur kalır, Kayıkçı seferine çıkar veya başarılı İsyanla görevden düşerse yeni Kaptanı hemen seçin',
'char_67':'6-7 kişide Uzakgören ve Kıyıçizen aynı kurulum setinde bulunamaz',
'char_density':'6–7\n4–5\n2',
'loyalty_6':'6\n1\n5',
'loyalty_15':'15\n5\n10',
'gusto':'Kaptan Gusto yok',
'black_seal':'Siyah Mühür',
}
for k,v in checks.items():
    if v not in txt: fail.append(f'missing rule marker: {k}')

chars=['Uzakgören','Kıyıçizen','Dümenkurdu','Canhalatı','Tahtakakan','Dipgören','Rüzgârkoklayan','Kırık Kürek','Üç Anahtar','Güvertebaşı','İskele Sıçanı','Karga Yuvası','Kazanbaşı','Fare Nazırı','Papağan Mütercimi','Fıçı Bekçisi','Kafiye Belası','Karayı Özleyen','Yastıkçı','Tahtaya Vuran']
for name in chars:
    if name not in txt: fail.append(f'character missing: {name}')

if fail:
    print('SONUC: FAIL')
    for x in fail: print('-',x)
    sys.exit(1)
print('SONUC: PASS')
print('CARD_HASH',h(CARDS))
print('BASELINE_HASH',h(BASE))
print('CHARACTERS',len(chars),'PASS')
print('RULEBOOK_MARKERS',len(checks),'PASS')
